# -*- coding: utf-8 -*-
"""Génère le questionnaire d'enquête (Word), format formulaire :
cases à cocher (qualitatif), cases numériques |__| (quantitatif),
sections en bandeaux, lignes de tableaux alternées orange/blanc."""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

HERE = os.path.dirname(os.path.abspath(__file__))

ORANGE_FONCE = "C55A11"
ORANGE_MOYEN = "F4B183"
ORANGE_CLAIR = "FDEADA"

GROUPE = ("Groupe 1 : Gado Giovanni Jocelyn, Alagbe Abdou Hamid, "
          "Sié Rachid Traoré, Cheikh Sadibou Ngom")

doc = Document()
doc.core_properties.author = "Groupe 1 - ISE3 ENSAE"
doc.core_properties.title = "Questionnaire d'enquête - Produits d'assurance vie"
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(3)


# ---------------------------------------------------------------- utilitaires
def shade(cell, color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_text(cell, text, bold=False, size=10.5, color=None, center=False,
                  italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string(color)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def alterner(tbl, start=1):
    """Alterne orange clair / blanc à partir de la ligne `start`."""
    for k, row in enumerate(tbl.rows[start:]):
        if k % 2 == 0:
            for cell in row.cells:
                shade(cell, ORANGE_CLAIR)


def bandeau(text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_text(cell, text, bold=True, size=11.5, color="FFFFFF")
    shade(cell, ORANGE_FONCE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def boxes(n):
    return "|" + "__|" * n


def question(code, text, consigne=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(f"{code}. ")
    run.bold = True
    p.add_run(text)
    if consigne:
        r = p.add_run(f"   [{consigne}]")
        r.italic = True
        r.font.size = Pt(9)
    return p


def q_quali(code, text, options, multi=False, autre=False, cols=2):
    question(code, text,
             "plusieurs réponses possibles" if multi else "une seule réponse")
    opts = list(options)
    if autre:
        opts.append("Autre, précisez : " + "." * 30)
    nrows = (len(opts) + cols - 1) // cols
    tbl = doc.add_table(rows=nrows, cols=cols)
    for k, opt in enumerate(opts):
        cell = tbl.rows[k % nrows].cells[k // nrows]
        set_cell_text(cell, "☐  " + opt, size=10.5)
    for row in tbl.rows:
        row.height = Cm(0.5)


def q_quanti(code, text, nboxes, unite="", virgule=0):
    p = question(code, text)
    p.add_run("      ")
    val = boxes(nboxes)
    if virgule:
        val += "," + boxes(virgule)[1:]
    run = p.add_run(val)
    run.bold = True
    if unite:
        p.add_run(f"  {unite}")


def q_date(code, text):
    p = question(code, text)
    p.add_run("      ")
    run = p.add_run("|__|__| / |__|__| / |__|__|__|__|")
    run.bold = True
    p.add_run("  (JJ/MM/AAAA)")


def q_libre(code, text, lignes=1):
    question(code, text, "réponse libre")
    for _ in range(lignes):
        p = doc.add_paragraph("." * 105)
        p.paragraph_format.space_after = Pt(2)


# ------------------------------------------------------------------- en-tête
titre = doc.add_table(rows=4, cols=1)
titre.style = "Table Grid"
set_cell_text(titre.rows[0].cells[0],
              "ENSAE Pierre Ndiaye - ISE 3 - Cours d'Assurance Vie",
              bold=True, size=11, center=True)
shade(titre.rows[0].cells[0], ORANGE_CLAIR)
set_cell_text(titre.rows[1].cells[0],
              "PROJET I : ENQUÊTE AUPRÈS DES INSTITUTIONS D'ASSURANCE",
              bold=True, size=14, center=True, color="FFFFFF")
shade(titre.rows[1].cells[0], ORANGE_FONCE)
set_cell_text(titre.rows[2].cells[0],
              "Questionnaire de recueil des produits d'assurance vie et de "
              "collecte des données de tarification",
              size=10.5, center=True, italic=True)
set_cell_text(titre.rows[3].cells[0],
              GROUPE + " - Année académique 2025-2026",
              size=9.5, center=True, italic=True)
doc.add_paragraph()

# fiche technique de l'entretien
ft = doc.add_table(rows=2, cols=4)
ft.style = "Table Grid"
for j, (lbl, val) in enumerate([
    ("N° du questionnaire", boxes(2)),
    ("Date de l'entretien", "|__|__|/|__|__|/|__|__|"),
    ("Heure de début", "|__|__| h |__|__|"),
    ("Durée (min)", boxes(3)),
]):
    set_cell_text(ft.rows[0].cells[j], lbl, bold=True, size=9, center=True)
    shade(ft.rows[0].cells[j], ORANGE_MOYEN)
    set_cell_text(ft.rows[1].cells[j], val, size=10, center=True)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
r = p.add_run("Mode de collecte :  ")
r.bold = True
p.add_run("☐ Face-à-face      ☐ Téléphone      ☐ En ligne")

# consignes : un tableau à une ligne par consigne (pas de retour à la ligne en cellule)
cons = doc.add_table(rows=4, cols=1)
cons.style = "Table Grid"
set_cell_text(cons.rows[0].cells[0], "CONSIGNES DE REMPLISSAGE",
              bold=True, size=9.5, center=True, color="FFFFFF")
shade(cons.rows[0].cells[0], ORANGE_FONCE)
for i, txt in enumerate([
    "☐  cocher la (ou les) case(s) correspondante(s), questions qualitatives",
    "|__|  inscrire un chiffre par case, questions quantitatives",
    "………  réponse libre (texte)",
], start=1):
    set_cell_text(cons.rows[i].cells[0], txt, size=9.5, italic=True)
alterner(cons)
p = doc.add_paragraph()
r = p.add_run("Confidentialité : les informations recueillies sont anonymisées "
              "et utilisées à des fins strictement pédagogiques.")
r.italic = True
r.font.size = Pt(9)

# ------------------------------------------------------------------- VOLET A
bandeau("VOLET A : IDENTIFICATION DE L'INSTITUTION")
q_libre("A1", "Dénomination de la compagnie :")
q_quali("A2", "Statut juridique :",
        ["Société anonyme (SA)", "Mutuelle d'assurance",
         "Filiale d'un groupe régional / international"], autre=True)
q_quanti("A3", "Année d'agrément au Sénégal (branche vie) :", 4)
q_quanti("A4", "Effectif total de la compagnie :", 4, "agents")
q_quali("A5", "Fonction de la personne interrogée :",
        ["Actuaire", "Directeur technique", "Souscripteur vie",
         "Responsable commercial"], autre=True)

# ------------------------------------------------------------------- VOLET B
bandeau("VOLET B : PORTEFEUILLE DE PRODUITS VIE")
q_quali("B1", "Quels produits d'assurance vie commercialisez-vous ?",
        ["Indemnités de Fin de Carrière (IFC)",
         "IFC Plus (décès du personnel, CCNI)",
         "Temporaire décès individuelle",
         "Décès emprunteur (crédit bancaire)",
         "Épargne / capitalisation",
         "Retraite complémentaire",
         "Rente éducation / rente de conjoint",
         "Mixte (épargne + décès)"],
        multi=True, autre=True)

question("B2", "Classez vos trois produits les plus commercialisés "
               "(exercice 2025) :")
tb = doc.add_table(rows=4, cols=3)
tb.style = "Table Grid"
for j, h in enumerate(["Rang", "Produit", "Primes émises 2025 (FCFA)"]):
    set_cell_text(tb.rows[0].cells[j], h, bold=True, size=9.5, center=True)
    shade(tb.rows[0].cells[j], ORANGE_MOYEN)
repeat_header(tb.rows[0])
for i in range(1, 4):
    set_cell_text(tb.rows[i].cells[0], f"n°{i}", center=True, size=10)
    set_cell_text(tb.rows[i].cells[2], boxes(12), center=True, size=10)
alterner(tb)
tb.columns[0].width = Cm(1.8)
tb.columns[1].width = Cm(8.5)
tb.columns[2].width = Cm(6.2)

q_quanti("B3", "Part du produit n°1 dans les primes vie de la compagnie :",
         2, "%", virgule=1)
q_quanti("B4", "Nombre de contrats du produit n°1 en portefeuille :", 5)
q_quali("B5", "Quelle clientèle porte principalement ce produit ?",
        ["Grandes entreprises", "PME / PMI", "Institutions publiques",
         "Particuliers"], multi=True)

# ------------------------------------------------------------------- VOLET C
bandeau("VOLET C : CARACTÉRISTIQUES TECHNIQUES DU PRODUIT LE PLUS COMMERCIALISÉ")
q_quali("C1", "Convention de référence pour le calcul des droits :",
        ["Convention Collective Nationale Interprofessionnelle (CCNI)",
         "Convention collective sectorielle"], autre=True, cols=1)
q_quali("C2", "Table de mortalité utilisée :",
        ["CIMA H", "CIMA F", "Table d'expérience"], autre=True)
q_quanti("C3", "Taux technique (max 3,5 %, art. 338 code CIMA) :", 1, "%",
         virgule=2)
q_quanti("C4", "Taux d'évolution des salaires (par défaut) :", 1, "%", virgule=2)
q_quanti("C5", "Taux de turn-over (par défaut) :", 1, "%", virgule=2)
q_quanti("C6", "Taux annuel de licenciement (par défaut) :", 1, "%", virgule=2)
q_quanti("C7", "Chargements de gestion (en % des primes) :", 2, "%", virgule=1)
q_quanti("C8", "Âge de départ à la retraite retenu (référence IPRES) :", 2, "ans")
q_quali("C9", "Modalités de financement proposées à l'entreprise :",
        ["Prime unique (engagement global)",
         "Dette actuarielle + charge normale annuelle",
         "Amortissement de la dette sur N exercices",
         "Versements libres"], multi=True, cols=1)
q_quanti("C9bis", "Si amortissement : durée proposée :", 2, "ans")
q_quali("C10", "Le fonds participe-t-il aux bénéfices ?", ["Oui", "Non"])
p = question("C10bis", "Si oui, à quelles hauteurs ?")
p.add_run("   financiers : ")
p.add_run(boxes(2) + " %").bold = True
p.add_run("      techniques : ")
p.add_run(boxes(2) + " %").bold = True
q_quali("C11", "Fréquence recommandée de réactualisation de l'étude :",
        ["Annuelle", "Triennale"], autre=True)

# ------------------------------------------------------------------- VOLET D
bandeau("VOLET D : FICHE DE COLLECTE DES DONNÉES DU PERSONNEL "
        "(entreprise souscriptrice)")
doc.add_paragraph(
    "Fiche remise à l'entreprise souhaitant souscrire le contrat. Les données "
    "ci-dessous sont demandées pour CHAQUE salarié ; elles constituent la base "
    "de données de l'évaluation actuarielle."
)
q_libre("D1", "Dénomination de l'entreprise :")
q_libre("D2", "Secteur d'activité :")
q_date("D3", "Date d'évaluation souhaitée (clôture d'exercice) :")
q_quanti("D4", "Effectif total à la date d'évaluation :", 4, "salariés")
q_quanti("D5", "Masse salariale annuelle brute :", 12, "FCFA")
q_quanti("D6", "Nombre de mois de salaire par an :", 2, "(12 ou 13)")
q_quali("D7", "Existe-t-il un fonds IFC déjà constitué ?", ["Oui", "Non"])
q_quanti("D7bis", "Si oui, montant du fonds :", 12, "FCFA")

question("D8", "Grille de recensement du personnel", "une ligne par salarié")
grid = doc.add_table(rows=6, cols=6)
grid.style = "Table Grid"
heads = ["Matricule", "Sexe (M/F)", "Catégorie*", "Date de naissance",
         "Date d'embauche", "Salaire mensuel brut (FCFA)"]
for j, h in enumerate(heads):
    set_cell_text(grid.rows[0].cells[j], h, bold=True, size=8.5, center=True)
    shade(grid.rows[0].cells[j], ORANGE_MOYEN)
repeat_header(grid.rows[0])
for i in range(1, 6):
    for j in range(6):
        grid.rows[i].cells[j].text = ""
    grid.rows[i].height = Cm(0.55)
alterner(grid)
p = doc.add_paragraph()
r = p.add_run("* Cadre dirigeant / Cadre / Agent de maîtrise / Employé / "
              "Ouvrier. Dates au format JJ/MM/AAAA ; salaire bonus compris. "
              "Joindre le fichier complet si l'effectif dépasse la grille.")
r.font.size = Pt(8.5)
r.italic = True

# ---------------------------------------------------------------------- fin
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Nous vous remercions de votre collaboration.")
r.bold = True
r.italic = True

out = os.path.join(HERE, "Questionnaire_enquete_IFC.docx")
doc.save(out)
print("Questionnaire écrit :", out)
